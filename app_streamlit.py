#!/usr/bin/env python3
# qc_app_single.py
# One-file Quality Reviewer + Streamlit UI
# Author: M365 Copilot for Jagamohan Padhy

import os
import re
import html
import json
import difflib
import tempfile
from dataclasses import dataclass, field
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime

import streamlit as st
import pandas as pd
from xml.etree import ElementTree as ET

# Optional imports for document reading
try:
    from PyPDF2 import PdfReader
except Exception:
    PdfReader = None

try:
    import docx  # python-docx
except Exception:
    docx = None

# ----------------------------- Utilities -----------------------------

def read_text_from_pdf(path: str) -> str:
    if PdfReader is None:
        raise RuntimeError("PyPDF2 not available to read PDFs.")
    reader = PdfReader(path)
    texts = []
    for page in reader.pages:
        try:
            texts.append(page.extract_text() or "")
        except Exception:
            texts.append("")
    return "\n".join(texts)

def read_text_from_docx(path: str) -> str:
    if docx is None:
        raise RuntimeError("python-docx not available to read DOCX.")
    d = docx.Document(path)
    parts = []
    for p in d.paragraphs:
        parts.append(p.text)
    for t in d.tables:
        for row in t.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)

def read_text_from_txt(path: str) -> str:
    with open(path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()

def read_csv_columns(path: str) -> pd.DataFrame:
    try:
        return pd.read_csv(path)
    except Exception:
        return pd.read_csv(path, encoding='latin-1')

def read_table_df(path: str) -> pd.DataFrame:
    ext = os.path.splitext(path)[1].lower()
    if ext == '.csv':
        return read_csv_columns(path)
    if ext == '.xlsx':
        return pd.read_excel(path, engine='openpyxl')
    if ext == '.xls':
        return pd.read_excel(path, engine='xlrd')
    raise RuntimeError(f'Unsupported tabular file type: {ext}')

def normalize_whitespace(s: str) -> str:
    return re.sub(r"\s+", " ", s or "").strip()

_punct_tbl = str.maketrans({c: "" for c in ('`~!@#$%^&*()-_=+[{]}\\|;:\'\",<.>/?')})

def simple_remove_punct(s: str) -> str:
    return (s or "").translate(_punct_tbl)

def alnum_only(s: str) -> str:
    return re.sub(r"[^0-9A-Za-z]", "", s or "")

def to_lower(s: str) -> str:
    return (s or "").lower()

def to_upper(s: str) -> str:
    return (s or "").upper()

COMMON_DATE_FORMATS = [
    "%Y-%m-%d", "%d-%m-%Y", "%d/%m/%Y", "%m/%d/%Y", "%d-%b-%Y", "%d %b %Y",
    "%Y/%m/%d", "%b %d, %Y", "%d.%m.%Y", "%Y.%m.%d",
]

def parse_date(s: str) -> Optional[datetime]:
    s = (s or "").strip()
    if not s:
        return None
    for fmt in COMMON_DATE_FORMATS:
        try:
            return datetime.strptime(s, fmt)
        except Exception:
            pass
    m = re.match(r"^(\d{4})(\d{2})(\d{2})$", s)
    if m:
        try:
            return datetime(int(m.group(1)), int(m.group(2)), int(m.group(3)))
        except Exception:
            return None
    return None

# ----------------------------- Config Models -----------------------------

@dataclass
class SourceSpec:
    method: str  # 'regex' | 'csv_column'
    pattern: Optional[str] = None
    flags: List[str] = field(default_factory=list)
    csv_column: Optional[str] = None

@dataclass
class XmlSpec:
    xpath: str

@dataclass
class ComparisonSpec:
    mode: str = "exact"  # exact | fuzzy | numeric | date | enum | bool
    tolerance: Optional[float] = None
    pct_tolerance: Optional[float] = None
    date_tolerance_days: Optional[int] = None
    fuzzy_threshold: float = 0.9

@dataclass
class FieldSpec:
    name: str
    type: str = "string"  # string, int, float, date, enum, bool
    source: SourceSpec = None
    xml: XmlSpec = None
    normalize: List[str] = field(default_factory=list)
    comparison: ComparisonSpec = field(default_factory=ComparisonSpec)
    allowed_values: List[str] = field(default_factory=list)  # for enum

# ----------------------------- Extraction -----------------------------

def apply_normalizers(val: Optional[str], normalizers: List[str]) -> Optional[str]:
    if val is None:
        return None
    s = val
    for n in normalizers:
        if n == 'strip':
            s = s.strip()
        elif n == 'lower':
            s = to_lower(s)
        elif n == 'upper':
            s = to_upper(s)
        elif n in ('collapse_whitespace', 'normalize_whitespace'):
            s = normalize_whitespace(s)
        elif n == 'alnum_only':
            s = alnum_only(s)
        elif n == 'remove_punctuation':
            s = simple_remove_punct(s)
    return s

def compile_flags(flags: List[str]) -> int:
    f = 0
    for fl in flags or []:
        u = fl.upper()
        if u == 'IGNORECASE':
            f |= re.IGNORECASE
        elif u == 'MULTILINE':
            f |= re.MULTILINE
        elif u == 'DOTALL':
            f |= re.DOTALL
    return f

def extract_from_source_text(text: str, spec: SourceSpec) -> Optional[str]:
    if spec.method == 'regex':
        if not spec.pattern:
            return None
        flags = compile_flags(spec.flags)
        m = re.search(spec.pattern, text, flags)
        if m:
            if m.lastindex:
                return m.group(1)
            return m.group(0)
        return None
    raise ValueError(f"Unsupported source method for text: {spec.method}")

def extract_from_source_csv(df: pd.DataFrame, spec: SourceSpec) -> Optional[str]:
    if spec.method == 'csv_column':
        if spec.csv_column and spec.csv_column in df.columns:
            series = df[spec.csv_column].dropna()
            if not series.empty:
                return str(series.iloc[0])
        return None
    raise ValueError(f"Unsupported source method for csv: {spec.method}")

def extract_from_xml(tree: ET.ElementTree, spec: XmlSpec) -> Optional[str]:
    try:
        node = tree.find(spec.xpath)
        if node is None:
            xp = spec.xpath
            if xp.endswith('/text()'):
                xp2 = xp[:-7]
                node = tree.find(xp2)
                if node is not None:
                    return (node.text or '').strip()
            return None
        return (node.text or '').strip()
    except Exception:
        return None

# ----------------------------- Comparison -----------------------------

def compare_values(field: FieldSpec, src_val: Optional[str], xml_val: Optional[str]) -> Tuple[str, Dict[str, Any]]:
    details: Dict[str, Any] = {
        'field': field.name,
        'type': field.type,
        'comparator': field.comparison.mode,
        'source_value': src_val,
        'xml_value': xml_val,
    }

    if src_val is None and xml_val is None:
        return 'NA', {**details, 'note': 'Both values missing'}
    if src_val is None or xml_val is None:
        return 'FAIL', {**details, 'note': 'One of the values is missing'}

    src_n = apply_normalizers(src_val, field.normalize)
    xml_n = apply_normalizers(xml_val, field.normalize)
    mode = (field.comparison.mode or 'exact').lower()

    if field.type in ('int', 'float') or mode == 'numeric':
        try:
            s_num = float(src_n)
            x_num = float(xml_n)
        except Exception:
            return 'FAIL', {**details, 'note': 'Numeric parsing failed'}
        tol = field.comparison.tolerance
        pct = field.comparison.pct_tolerance
        if tol is not None:
            diff = abs(s_num - x_num)
            ok = diff <= tol
            return ('PASS' if ok else 'FAIL', {**details, 'diff': diff, 'tolerance': tol})
        if pct is not None:
            base = abs(s_num) if s_num != 0 else 1.0
            diff_pct = abs(s_num - x_num) / base
            ok = diff_pct <= pct
            return ('PASS' if ok else 'FAIL', {**details, 'diff_pct': diff_pct, 'pct_tolerance': pct})
        ok = s_num == x_num
        return ('PASS' if ok else 'FAIL', {**details})

    if field.type == 'date' or mode == 'date':
        d1 = parse_date(src_n)
        d2 = parse_date(xml_n)
        if not d1 or not d2:
            return 'FAIL', {**details, 'note': 'Date parsing failed'}
        dtol = field.comparison.date_tolerance_days
        if dtol is not None:
            diff_days = abs((d1 - d2).days)
            ok = diff_days <= dtol
            return ('PASS' if ok else 'FAIL', {**details, 'diff_days': diff_days, 'tolerance_days': dtol})
        ok = d1.date() == d2.date()
        return ('PASS' if ok else 'FAIL', {**details})

    if field.type == 'bool':
        def to_bool(s: str) -> Optional[bool]:
            sl = (s or '').strip().lower()
            if sl in ('true','yes','y','1'): return True
            if sl in ('false','no','n','0'): return False
            return None
        b1 = to_bool(src_n)
        b2 = to_bool(xml_n)
        if b1 is None or b2 is None:
            return 'FAIL', {**details, 'note': 'Boolean parsing failed'}
        ok = b1 == b2
        return ('PASS' if ok else 'FAIL', {**details})

    if field.type == 'enum':
        if field.allowed_values:
            if src_n not in field.allowed_values:
                details['note'] = 'Source value not in allowed enum'
            if xml_n not in field.allowed_values:
                details['note'] = (details.get('note','')+'; ' if details.get('note') else '') + 'XML value not in allowed enum'
        ok = src_n == xml_n
        return ('PASS' if ok else 'FAIL', {**details})

    if mode == 'fuzzy':
        ratio = difflib.SequenceMatcher(a=src_n, b=xml_n).ratio()
        thr = field.comparison.fuzzy_threshold or 0.9
        ok = ratio >= thr
        return ('PASS' if ok else 'FAIL', {**details, 'similarity': ratio, 'threshold': thr})

    ok = src_n == xml_n
    return ('PASS' if ok else 'FAIL', {**details})

# ----------------------------- Engine (run_qc) -----------------------------

def load_config(path: str) -> Dict[str, Any]:
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)

def build_field_specs(cfg: Dict[str, Any]) -> List[FieldSpec]:
    fields_cfg = cfg.get('fields', [])
    result: List[FieldSpec] = []
    for f in fields_cfg:
        source_cfg = f.get('source', {})
        xml_cfg = f.get('xml', {})
        comp_cfg = f.get('comparison', {})
        fs = FieldSpec(
            name=f.get('name'),
            type=f.get('type', 'string'),
            source=SourceSpec(method=source_cfg.get('method'),
                              pattern=source_cfg.get('pattern'),
                              flags=source_cfg.get('flags', []),
                              csv_column=source_cfg.get('csv_column')),
            xml=XmlSpec(xpath=xml_cfg.get('xpath')),
            normalize=f.get('normalize', []),
            comparison=ComparisonSpec(
                mode=comp_cfg.get('mode', 'exact'),
                tolerance=comp_cfg.get('tolerance'),
                pct_tolerance=comp_cfg.get('pct_tolerance'),
                date_tolerance_days=comp_cfg.get('date_tolerance_days'),
                fuzzy_threshold=comp_cfg.get('fuzzy_threshold', 0.9),
            ),
            allowed_values=f.get('allowed_values', [])
        )
        result.append(fs)
    return result

def extract_source_values(source_path: str, fields: List[FieldSpec]) -> Dict[str, Optional[str]]:
    ext = os.path.splitext(source_path)[1].lower()
    text_cache: Optional[str] = None
    table_df: Optional[pd.DataFrame] = None

    def ensure_text():
        nonlocal text_cache
        if text_cache is not None:
            return text_cache
        if ext == '.pdf':
            text_cache = read_text_from_pdf(source_path)
        elif ext == '.docx':
            text_cache = read_text_from_docx(source_path)
        elif ext in ('.txt', '.text'):
            text_cache = read_text_from_txt(source_path)
        else:
            raise RuntimeError(f"Unsupported source file type for text extraction: {ext}")
        return text_cache

    def ensure_table():
        nonlocal table_df
        if table_df is not None:
            return table_df
        table_df = read_table_df(source_path)
        return table_df

    out: Dict[str, Optional[str]] = {}
    for f in fields:
        try:
            if f.source.method == 'regex':
                txt = ensure_text()
                out[f.name] = extract_from_source_text(txt, f.source)
            elif f.source.method == 'csv_column':
                if ext not in ('.csv', '.xlsx', '.xls'):
                    raise RuntimeError("Configured for csv_column but source is not a CSV/Excel file")
                df = ensure_table()
                out[f.name] = extract_from_source_csv(df, f.source)
            else:
                out[f.name] = None
        except Exception:
            out[f.name] = None
    return out

def extract_xml_values(xml_path: str, fields: List[FieldSpec]) -> Dict[str, Optional[str]]:
    tree = ET.parse(xml_path)
    out: Dict[str, Optional[str]] = {}
    for f in fields:
        try:
            out[f.name] = extract_from_xml(tree, f.xml)
        except Exception:
            out[f.name] = None
    return out

def generate_reports(results: List[Dict[str, Any]], outdir: str, base_name: str, formats: List[str], html_summary: bool = False) -> Dict[str, str]:
    os.makedirs(outdir, exist_ok=True)
    df = pd.DataFrame(results)
    paths: Dict[str, str] = {}

    if 'csv' in formats or 'all' in formats:
        p = os.path.join(outdir, f"{base_name}_qc_report.csv")
        df.to_csv(p, index=False)
        paths['csv'] = p
    if 'xlsx' in formats or 'all' in formats:
        p = os.path.join(outdir, f"{base_name}_qc_report.xlsx")
        with pd.ExcelWriter(p, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='QC Report')
        paths['xlsx'] = p
    if 'json' in formats or 'all' in formats:
        p = os.path.join(outdir, f"{base_name}_qc_report.json")
        df.to_json(p, orient='records', indent=2)
        paths['json'] = p

    if html_summary:
        p = os.path.join(outdir, f"{base_name}_qc_summary.html")
        passes = sum(1 for r in results if r.get('status') == 'PASS')
        fails = sum(1 for r in results if r.get('status') == 'FAIL')
        warns = sum(1 for r in results if r.get('status') == 'WARN')
        nas = sum(1 for r in results if r.get('status') == 'NA')
        total = len(results)
        table_rows = []
        for r in results:
            stt = r.get('status')
            color = {'PASS':'#d1fae5','FAIL':'#fee2e2','WARN':'#fef9c3','NA':'#e5e7eb'}.get(stt,'#ffffff')
            row = (
                f"<tr style='background:{color}'>"
                f"<td>{html.escape(str(r.get('field')))}</td>"
                f"<td>{html.escape(str(r.get('source_value')))}</td>"
                f"<td>{html.escape(str(r.get('xml_value')))}</td>"
                f"<td>{html.escape(str(stt))}</td>"
                f"<td>{html.escape(str(r.get('comparator')))}</td>"
                f"<td>{html.escape(str(r.get('note', '')))}</td>"
                f"</tr>"
            )
            table_rows.append(row)
        html_doc = f"""
        <html><head><meta charset='utf-8'><title>QC Summary</title>
        <style>body{{font-family:Arial, sans-serif}} table{{border-collapse:collapse;width:100%}} td,th{{border:1px solid #ddd;padding:8px}}</style>
        </head><body>
        <h2>Quality Review Summary</h2>
        <p><b>Total:</b> {total} | <b>PASS:</b> {passes} | <b>FAIL:</b> {fails} | <b>WARN:</b> {warns} | <b>NA:</b> {nas}</p>
        <table><thead><tr><th>Field</th><th>Source</th><th>XML</th><th>Status</th><th>Comparator</th><th>Notes</th></tr></thead>
        <tbody>
        {''.join(table_rows)}
        </tbody></table>
        </body></html>
        """
        with open(p, 'w', encoding='utf-8') as f:
            f.write(html_doc)
        paths['html'] = p

    return paths

def run_qc(source_path: str, xml_path: str, config_path: str, outdir: str, report_format: str, html_summary: bool) -> Dict[str, str]:
    cfg = load_config(config_path)
    fields = build_field_specs(cfg)

    src_vals = extract_source_values(source_path, fields)
    xml_vals = extract_xml_values(xml_path, fields)

    results: List[Dict[str, Any]] = []
    for f in fields:
        s_val = src_vals.get(f.name)
        x_val = xml_vals.get(f.name)
        status, detail = compare_values(f, s_val, x_val)
        results.append({**detail, 'status': status})

    base = os.path.splitext(os.path.basename(source_path))[0] + "__" + os.path.splitext(os.path.basename(xml_path))[0]
    fmts = ['csv'] if report_format == 'csv' else ['xlsx'] if report_format == 'xlsx' else ['json'] if report_format == 'json' else ['all']
    paths = generate_reports(results, outdir, base, fmts, html_summary)
    return paths

# ----------------------------- Streamlit UI -----------------------------

st.set_page_config(page_title="Quality Reviewer", layout="wide")
st.title("Quality Reviewer (Source vs Processed XML)")
st.caption("Upload a source document and a processed XML, provide a config, and generate a QC report.")

with st.sidebar:
    st.header("Settings")
    report_format = st.selectbox("Report format", ["all", "xlsx", "csv", "json"], index=0)
    gen_html = st.checkbox("Generate HTML summary", value=True)
    outdir_name = st.text_input("Output folder name", value=f"qc_output_{datetime.now().strftime('%Y%m%d_%H%M%S')}")

col1, col2, col3 = st.columns(3)
source_file = col1.file_uploader("Source (PDF/DOCX/TXT/CSV/XLSX/XLS)", type=["pdf","docx","txt","csv","xlsx","xls"])
xml_file = col2.file_uploader("Processed XML", type=["xml"])
config_file = col3.file_uploader("Config (JSON)", type=["json"])

def _infer_mime(path: str) -> str:
    p = path.lower()
    if p.endswith(".csv"):
        return "text/csv"
    if p.endswith(".json"):
        return "application/json"
    if p.endswith(".xlsx"):
        return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    if p.endswith(".html"):
        return "text/html"
    return "application/octet-stream"

def _preview_config(cfg_bytes: bytes):
    try:
        txt = cfg_bytes.decode("utf-8", errors="ignore")
        obj = json.loads(txt)
        st.markdown("**Config (JSON) preview:**")
        pretty = json.dumps(obj, indent=2)
        st.code(pretty[:2000] + ("..." if len(pretty) > 2000 else ""), language="json")
    except Exception:
        pass

if st.button("Run Quality Check", use_container_width=True):
    if not (source_file and xml_file and config_file):
        st.error("Please upload Source, XML, and Config files.")
        st.stop()

    with st.expander("Uploaded files"):
        st.write({
            "source": source_file.name if source_file else None,
            "xml": xml_file.name if xml_file else None,
            "config": config_file.name if config_file else None
        })
        _preview_config(config_file.getvalue())

    with tempfile.TemporaryDirectory() as tmpdir:
        # Save uploads
        src_path = os.path.join(tmpdir, source_file.name)
        with open(src_path, "wb") as f:
            f.write(source_file.read())
        xml_path = os.path.join(tmpdir, xml_file.name)
        with open(xml_path, "wb") as f:
            f.write(xml_file.read())
        cfg_path = os.path.join(tmpdir, config_file.name)
        with open(cfg_path, "wb") as f:
            f.write(config_file.read())

        outdir = os.path.join(tmpdir, outdir_name)

        with st.spinner("Running QC..."):
            try:
                paths = run_qc(src_path, xml_path, cfg_path, outdir, report_format, gen_html)
            except Exception as e:
                st.error("QC failed. See details below:")
                st.exception(e)
                st.stop()

        st.success("QC complete. Download your outputs below.")

        # Download buttons
        if not paths:
            st.warning("No outputs were generated. Please check your config/inputs.")
        else:
            for key, p in paths.items():
                label = f"Download {key.upper()} report"
                with open(p, "rb") as fh:
                    st.download_button(
                        label=label,
                        data=fh.read(),
                        file_name=os.path.basename(p),
                        mime=_infer_mime(p),
                        use_container_width=True
                    )

        # Preview
        if "csv" in paths and paths["csv"].endswith(".csv"):
            try:
                df = pd.read_csv(paths["csv"])
                st.subheader("QC Report Preview")
                st.dataframe(df, use_container_width=True)
            except Exception:
                pass
        elif "json" in paths and paths["json"].endswith(".json"):
            try:
                df = pd.read_json(paths["json"])
                st.subheader("QC Report Preview")
                st.dataframe(df, use_container_width=True)
            except Exception:
                pass

st.caption("Tip: For scanned PDFs (no text layer), run OCR first or provide DOCX/TXT/CSV/XLSX.")